from __future__ import annotations

import re
from pathlib import Path
from typing import Any

import fitz
from docx import Document

from . import conversion_engine as legacy
from .processing_quality import DEFAULT_OCR_DPI, _ocr_image_model, _render_page, _searchable_from_images, _write_text


def selected_pdf_pages(raw: object, total: int) -> list[int]:
    value = str(raw or "all").strip().lower()
    if value in {"", "all", "*"}:
        return list(range(1, total + 1))
    selected: set[int] = set()
    for item in value.split(","):
        token = item.strip()
        if not token:
            continue
        match = re.fullmatch(r"(\d+)(?:\s*-\s*(\d+))?", token)
        if not match:
            raise ValueError("OCR pages must look like all, 2, 1-3, or 1,4,7-9.")
        start = int(match.group(1))
        end = int(match.group(2) or start)
        if start < 1 or end < start or end > total:
            raise ValueError(f"OCR page range {token} is outside 1-{total}.")
        selected.update(range(start, end + 1))
    if not selected:
        raise ValueError("Select at least one OCR page.")
    return sorted(selected)


def _models(source: Path, options: dict[str, Any]) -> list[dict[str, Any]]:
    dpi = max(150, min(400, int(options.get("dpi", DEFAULT_OCR_DPI))))
    models: list[dict[str, Any]] = []
    with fitz.open(source) as document:
        if document.page_count < 1:
            raise ValueError("The PDF has no pages.")
        if document.page_count > legacy.MAX_PDF_PAGES:
            raise ValueError(f"This PDF has too many pages for one OCR job. Maximum: {legacy.MAX_PDF_PAGES}.")
        for page_number in selected_pdf_pages(options.get("pages", "all"), document.page_count):
            image = _render_page(document[page_number - 1], dpi)
            try:
                model = _ocr_image_model(image, options)
            finally:
                image.close()
            models.append({**model, "page": page_number})
    return models


def _word(models: list[dict[str, Any]], output: Path, title: str) -> None:
    if not models:
        raise ValueError("No OCR pages were selected.")
    document = Document()
    document.core_properties.title = title
    for index, model in enumerate(models):
        if index:
            document.add_page_break()
        document.add_heading(f"Page {model['page']}", level=2)
        text = str(model.get("text") or "").strip()
        if not text:
            document.add_paragraph("[No readable OCR text detected on this page]")
            continue
        for paragraph in re.split(r"\n\s*\n", text):
            if paragraph.strip():
                document.add_paragraph(paragraph.strip())
    document.save(output)


def _searchable(source: Path, output: Path, options: dict[str, Any], workdir: Path) -> None:
    dpi = max(150, min(400, int(options.get("dpi", DEFAULT_OCR_DPI))))
    rendered: list[Path] = []
    try:
        with fitz.open(source) as document:
            if document.page_count < 1:
                raise ValueError("The PDF has no pages.")
            for page_number in selected_pdf_pages(options.get("pages", "all"), document.page_count):
                image = _render_page(document[page_number - 1], dpi)
                path = workdir / f"ocr-selected-page-{page_number:03d}.png"
                try:
                    image.save(path, "PNG")
                finally:
                    image.close()
                rendered.append(path)
        _searchable_from_images(rendered, output, options, workdir)
    finally:
        for path in rendered:
            path.unlink(missing_ok=True)


def run_selected_pdf_ocr(
    processor: str,
    source: Path,
    output: Path,
    options: dict[str, Any],
    workdir: Path,
) -> None:
    if processor == "ocr_pdf_text":
        _write_text(_models(source, options), output)
        return
    if processor == "ocr_pdf_word":
        _word(_models(source, options), output, source.stem)
        return
    if processor == "ocr_pdf_searchable":
        _searchable(source, output, options, workdir)
        return
    raise ValueError(f"Unsupported selected-page OCR processor: {processor}")
