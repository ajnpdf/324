from __future__ import annotations
import csv
import io
import json
import math
import re
import statistics
import zipfile
from pathlib import Path
from typing import Any, Iterable
import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from openpyxl import Workbook
from PIL import Image, ImageEnhance, ImageFilter, ImageOps, ImageSequence
from . import conversion_engine as legacy
from .r19_fidelity import run_r19_conversion
MIN_NATIVE_TEXT_CHARS = 24
MIN_USEFUL_recognition_CHARS = 3
DEFAULT_SCAN_DPI = 240
MAX_TABLES_PER_DOCUMENT = 64
MAX_DOCX_IMAGES = 64

def _projection_score(image: Image.Image) -> float:
    gray = ImageOps.grayscale(image)
    if gray.width > 1000:
        ratio = 1000 / gray.width
        gray = gray.resize((1000, max(1, int(gray.height * ratio))), Image.Resampling.BILINEAR)
    gray = ImageOps.autocontrast(gray)
    pixels = gray.load()
    row_scores: list[float] = []
    for y in range(gray.height):
        darkness = 0
        for x in range(gray.width):
            darkness += 255 - int(pixels[x, y])
        row_scores.append(float(darkness))
    if len(row_scores) < 2:
        return 0.0
    return statistics.pvariance(row_scores)

def _deskew(image: Image.Image) -> Image.Image:
    if min(image.size) < 120:
        return image
    best_angle = 0.0
    best_score = _projection_score(image)
    for angle in (-4, -3, -2, -1, 1, 2, 3, 4):
        candidate = image.rotate(angle, expand=True, fillcolor='white')
        score = _projection_score(candidate)
        candidate.close()
        if score > best_score * 1.015:
            best_score = score
            best_angle = float(angle)
    if not best_angle:
        return image
    return image.rotate(best_angle, expand=True, fillcolor='white')

def _orient_with_osd(image: Image.Image) -> Image.Image:
    try:
        info = retired_engine.image_to_osd(image, output_type=Output.DICT)
        rotate = int(info.get('rotate') or 0) % 360
        confidence = float(info.get('orientation_conf') or 0)
        if rotate and confidence >= 3:
            return image.rotate(-rotate, expand=True, fillcolor='white')
    except Exception:
        pass
    return image

def _prepare_scan_image(image: Image.Image, options: dict[str, Any], *, preserve_color: bool=False) -> Image.Image:
    prepared = ImageOps.exif_transpose(image.copy())
    if bool(options.get('auto_rotate', True)):
        rotated = _orient_with_osd(prepared)
        if rotated is not prepared:
            prepared.close()
            prepared = rotated
    if bool(options.get('deskew', True)):
        skewed = _deskew(prepared)
        if skewed is not prepared:
            prepared.close()
            prepared = skewed
    if preserve_color:
        return ImageOps.autocontrast(prepared.convert('RGB'), cutoff=0.5)
    gray = ImageOps.grayscale(prepared)
    prepared.close()
    gray = ImageOps.autocontrast(gray, cutoff=1)
    gray = ImageEnhance.Contrast(gray).enhance(float(options.get('contrast', 1.35)))
    if bool(options.get('denoise', True)):
        gray = gray.filter(ImageFilter.MedianFilter(size=3))
    gray = gray.filter(ImageFilter.SHARPEN)
    return gray.convert('RGB')

def _native_page_model(page: fitz.Page) -> dict[str, Any]:
    raw = page.get_text('dict', sort=True)
    blocks: list[dict[str, Any]] = []
    text_lines: list[str] = []
    font_sizes: list[float] = []
    for block in raw.get('blocks', []):
        if int(block.get('type', 0)) != 0:
            continue
        block_lines: list[dict[str, Any]] = []
        for line in block.get('lines', []):
            spans_out: list[dict[str, Any]] = []
            line_text: list[str] = []
            for span in line.get('spans', []):
                text = str(span.get('text') or '')
                if not text:
                    continue
                size = float(span.get('size') or 11)
                font_sizes.append(size)
                spans_out.append({'text': text, 'size': size, 'font': str(span.get('font') or ''), 'flags': int(span.get('flags') or 0), 'bbox': tuple(span.get('bbox') or (0, 0, 0, 0))})
                line_text.append(text)
            joined = ''.join(line_text).strip()
            if joined:
                text_lines.append(joined)
                block_lines.append({'text': joined, 'spans': spans_out, 'bbox': tuple(line.get('bbox') or (0, 0, 0, 0))})
        if block_lines:
            blocks.append({'lines': block_lines, 'bbox': tuple(block.get('bbox') or (0, 0, 0, 0))})
    text = ''.join(text_lines).strip()
    return {'text': text, 'lines': text_lines, 'blocks': blocks, 'font_sizes': font_sizes, 'page_width': float(page.rect.width), 'page_height': float(page.rect.height), 'source': 'native'}

def _render_page(page: fitz.Page, dpi: int) -> Image.Image:
    matrix = fitz.Matrix(dpi / 72.0, dpi / 72.0)
    pix = page.get_pixmap(matrix=matrix, alpha=False)
    return Image.frombytes('RGB', (pix.width, pix.height), pix.samples)

def _extract_pdf_models(source: Path, options: dict[str, Any], *, force_scan: bool=False) -> list[dict[str, Any]]:
    dpi = max(150, min(400, int(options.get('dpi', DEFAULT_SCAN_DPI))))
    models: list[dict[str, Any]] = []
    with fitz.open(source) as document:
        if document.page_count < 1:
            raise ValueError('The PDF has no pages.')
        if document.page_count > legacy.MAX_PDF_PAGES:
            raise ValueError(f'This PDF has too many pages for one processing job. Maximum:{legacy.MAX_PDF_PAGES}.')
        for index, page in enumerate(document, start=1):
            native = _native_page_model(page)
            useful_native = len(re.sub('\\s+', '', native['text'])) >= MIN_NATIVE_TEXT_CHARS
            if force_scan or not useful_native:
                model = {**native, 'page': index}
            else:
                model = {**native, 'page': index, 'confidence': 100.0}
            models.append(model)
    return models

def _write_text(models: list[dict[str, Any]], output: Path) -> None:
    useful = [model for model in models if str(model.get('text') or '').strip()]
    if not useful:
        raise ValueError('No readable text was found in this PDF. Try with the correct language selected.')
    output.write_text(''.join((f"--- Page{model['page']}---{str(model.get('text') or '').strip()}" for model in models)), encoding='utf-8')

def _docx_add_native_block(document: Document, block: dict[str, Any], body_size: float, page_width: float) -> None:
    lines = block.get('lines') or []
    spans = [span for line in lines for span in line.get('spans', [])]
    if not spans:
        return
    average_size = statistics.mean((float(span.get('size') or body_size) for span in spans))
    style = None
    if body_size > 0 and average_size >= body_size * 1.65:
        style = 'Heading 1'
    elif body_size > 0 and average_size >= body_size * 1.3:
        style = 'Heading 2'
    paragraph = document.add_paragraph(style=style)
    bbox = block.get('bbox') or (0, 0, page_width, 0)
    center = (float(bbox[0]) + float(bbox[2])) / 2
    width = max(1.0, float(bbox[2]) - float(bbox[0]))
    if abs(center - page_width / 2) < page_width * 0.08 and width < page_width * 0.8:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    first_line = True
    for line in lines:
        if not first_line:
            paragraph.add_run().add_break()
        first_line = False
        for span in line.get('spans', []):
            run = paragraph.add_run(str(span.get('text') or ''))
            font = str(span.get('font') or '').lower()
            flags = int(span.get('flags') or 0)
            run.bold = bool(flags & 16) or 'bold' in font
            run.italic = bool(flags & 2) or 'italic' in font or 'oblique' in font
            size = max(7.0, min(36.0, float(span.get('size') or body_size)))
            run.font.size = Pt(size)

def _write_layout_docx(source: Path, models: list[dict[str, Any]], output: Path, options: dict[str, Any]) -> None:
    document = Document()
    document.core_properties.title = source.stem
    image_count = 0
    with fitz.open(source) as pdf:
        all_sizes = [size for model in models for size in model.get('font_sizes', []) if 5 <= float(size) <= 36]
        body_size = statistics.median(all_sizes) if all_sizes else 11.0
        for page_index, model in enumerate(models):
            if page_index:
                document.add_page_break()
            if model.get('source') == 'native' and model.get('blocks'):
                for block in model['blocks']:
                    _docx_add_native_block(document, block, body_size, float(model.get('page_width') or 612))
            else:
                text = str(model.get('text') or '').strip()
                if text:
                    for paragraph_text in re.split('\\n\\s*\\n', text):
                        paragraph_text = paragraph_text.strip()
                        if paragraph_text:
                            document.add_paragraph(paragraph_text)
            if bool(options.get('include_images', True)) and image_count < MAX_DOCX_IMAGES:
                page = pdf[page_index]
                for image_info in page.get_images(full=True):
                    if image_count >= MAX_DOCX_IMAGES:
                        break
                    try:
                        extracted = pdf.extract_image(image_info[0])
                        data = extracted.get('image')
                        if not data or len(data) < 4096:
                            continue
                        with Image.open(io.BytesIO(data)) as probe:
                            if probe.width < 80 or probe.height < 80:
                                continue
                        document.add_picture(io.BytesIO(data), width=Inches(5.8))
                        image_count += 1
                    except Exception:
                        continue
    document.save(output)

def _table_rows(table: Any) -> list[list[str]]:
    rows = table.extract()
    cleaned: list[list[str]] = []
    for row in rows or []:
        values = [str(value or '').strip() for value in row]
        if any(values):
            cleaned.append(values)
    return cleaned

def _pdf_tables(source: Path) -> list[tuple[int, int, list[list[str]]]]:
    found: list[tuple[int, int, list[list[str]]]] = []
    with fitz.open(source) as document:
        for page_index, page in enumerate(document, start=1):
            try:
                finder = page.find_tables()
                tables = list(getattr(finder, 'tables', []) or [])
            except Exception:
                tables = []
            for table_index, table in enumerate(tables, start=1):
                rows = _table_rows(table)
                if rows and max((len(row) for row in rows), default=0) >= 2:
                    found.append((page_index, table_index, rows))
                    if len(found) >= MAX_TABLES_PER_DOCUMENT:
                        return found
    return found

def _write_xlsx_tables(source: Path, output: Path, options: dict[str, Any]) -> None:
    tables = _pdf_tables(source)
    if not tables:
        if bool(options.get('allow_unstructured', False)):
            models = _extract_pdf_models(source, options)
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = 'PDF text'
            sheet.append(['Page', 'Line', 'Text'])
            for model in models:
                for line_number, line in enumerate(model.get('lines') or [], start=1):
                    sheet.append([model['page'], line_number, line])
            workbook.save(output)
            return
        raise ValueError('No structured table was detected in this PDF. Try PDF to Word or PDF to Text for this document.')
    workbook = Workbook()
    workbook.remove(workbook.active)
    for table_number, (page_number, page_table_number, rows) in enumerate(tables, start=1):
        name = f'P{page_number}-Table{page_table_number}'[:31]
        if name in workbook.sheetnames:
            name = f'Table{table_number}'[:31]
        sheet = workbook.create_sheet(name)
        for row in rows:
            sheet.append(row)
        for column_cells in sheet.columns:
            letter = column_cells[0].column_letter
            width = min(60, max(10, max((len(str(cell.value or '')) for cell in column_cells)) + 2))
            sheet.column_dimensions[letter].width = width
    workbook.save(output)

def _write_csv_tables(source: Path, output: Path, options: dict[str, Any]) -> None:
    tables = _pdf_tables(source)
    if not tables:
        raise ValueError('No structured table was detected in this PDF. Try PDF to Word or PDF to Text for this document.')
    with output.open('w', newline='', encoding='utf-8-sig') as handle:
        writer = csv.writer(handle)
        for index, (page_number, table_number, rows) in enumerate(tables):
            if index:
                writer.writerow([])
            writer.writerow([f'Page{page_number}- Table{table_number}'])
            writer.writerows(rows)

def _iter_input_frames(path: Path) -> Iterable[Image.Image]:
    with Image.open(path) as source:
        for frame_index, frame in enumerate(ImageSequence.Iterator(source), start=1):
            if frame_index > legacy.MAX_IMAGE_FRAMES:
                raise ValueError(f'This image contains too many frames. Maximum:{legacy.MAX_IMAGE_FRAMES}.')
            prepared = ImageOps.exif_transpose(frame.copy())
            if prepared.width * prepared.height > legacy.MAX_IMAGE_PIXELS:
                prepared.close()
                raise ValueError('The image dimensions are too large to process safely.')
            yield prepared

def _page_size_points(image: Image.Image, options: dict[str, Any]) -> tuple[float, float]:
    requested = str(options.get('page_size', 'auto')).lower()
    dpi = max(72.0, min(600.0, float(options.get('dpi', 150))))
    if requested == 'a4':
        width, height = (595.276, 841.89)
    elif requested in {'letter', 'us-letter'}:
        width, height = (612.0, 792.0)
    else:
        width = image.width * 72.0 / dpi
        height = image.height * 72.0 / dpi
    orientation = str(options.get('orientation', 'auto')).lower()
    if orientation == 'landscape' and height > width:
        width, height = (height, width)
    elif orientation == 'portrait' and width > height:
        width, height = (height, width)
    return (max(36.0, width), max(36.0, height))

def _images_to_pdf_quality(files: list[Path], output: Path, options: dict[str, Any], *, scan: bool=False) -> None:
    pdf = fitz.open()
    total_pixels = 0
    try:
        for file in files:
            for frame in _iter_input_frames(file):
                try:
                    total_pixels += frame.width * frame.height
                    if total_pixels > legacy.MAX_BATCH_PIXELS:
                        raise ValueError('The selected images are too large for one PDF job. Use fewer or smaller images.')
                    image = _prepare_scan_image(frame, options, preserve_color=not bool(options.get('grayscale', scan))) if scan else frame.convert('RGB')
                    try:
                        page_width, page_height = _page_size_points(image, options)
                        page = pdf.new_page(width=page_width, height=page_height)
                        margin_mm = max(0.0, min(50.0, float(options.get('margin_mm', 0))))
                        margin = margin_mm * 72.0 / 25.4
                        available = fitz.Rect(margin, margin, page_width - margin, page_height - margin)
                        if available.width <= 1 or available.height <= 1:
                            raise ValueError('The selected page margin leaves no room for the image.')
                        scale = min(available.width / image.width, available.height / image.height)
                        draw_w = image.width * scale
                        draw_h = image.height * scale
                        x0 = available.x0 + (available.width - draw_w) / 2
                        y0 = available.y0 + (available.height - draw_h) / 2
                        target = fitz.Rect(x0, y0, x0 + draw_w, y0 + draw_h)
                        buffer = io.BytesIO()
                        image.save(buffer, 'JPEG', quality=max(50, min(100, int(options.get('quality', 92)))), optimize=True)
                        page.insert_image(target, stream=buffer.getvalue(), keep_proportion=True)
                    finally:
                        if image is not frame:
                            image.close()
                finally:
                    frame.close()
        if pdf.page_count < 1:
            raise ValueError('No readable images were found.')
        pdf.save(output, garbage=4, deflate=True)
    finally:
        pdf.close()

def _selected_pages(options: dict[str, Any], page_count: int) -> list[int]:
    raw = options.get('pages')
    if raw in (None, '', [], 'all'):
        return list(range(page_count))
    selected: set[int] = set()
    items = raw if isinstance(raw, list) else str(raw).split(',')
    for item in items:
        if isinstance(item, int):
            selected.add(item - 1)
            continue
        token = str(item).strip()
        if not token:
            continue
        if '-' in token:
            start_raw, end_raw = token.split('-', 1)
            start, end = (int(start_raw), int(end_raw))
            for page in range(min(start, end), max(start, end) + 1):
                selected.add(page - 1)
        else:
            selected.add(int(token) - 1)
    valid = sorted((index for index in selected if 0 <= index < page_count))
    if not valid:
        raise ValueError('The selected page range does not contain any valid PDF pages.')
    return valid

def _pdf_to_images_quality(source: Path, output: Path, fmt: str, options: dict[str, Any], workdir: Path) -> None:
    dpi = max(72, min(600, int(options.get('dpi', 150))))
    quality = max(35, min(100, int(options.get('quality', 90))))
    generated: list[Path] = []
    with fitz.open(source) as document:
        selected = _selected_pages(options, document.page_count)
        total_pixels = 0
        images: list[Image.Image] = []
        try:
            for output_index, page_index in enumerate(selected, start=1):
                page = document[page_index]
                image = _render_page(page, dpi)
                total_pixels += image.width * image.height
                if total_pixels > legacy.MAX_RENDER_PIXELS:
                    image.close()
                    raise ValueError('The selected page count and resolution are too large for one job. Lower the resolution or choose fewer pages.')
                if fmt in {'gif', 'tiff'}:
                    images.append(image)
                    continue
                if fmt == 'svg':
                    image.close()
                    path = workdir / f'page-{page_index + 1:03d}.svg'
                    path.write_text(page.get_svg_image(), encoding='utf-8')
                    generated.append(path)
                    continue
                normalized = 'JPEG' if fmt in {'jpg', 'jpeg'} else 'HEIF' if fmt == 'heic' else fmt.upper()
                suffix = 'jpg' if fmt in {'jpg', 'jpeg'} else fmt
                path = workdir / f'page-{page_index + 1:03d}.{suffix}'
                try:
                    kwargs: dict[str, Any] = {}
                    if normalized in {'JPEG', 'WEBP', 'HEIF', 'AVIF'}:
                        kwargs['quality'] = quality
                    image.save(path, normalized, **kwargs)
                except Exception as exc:
                    raise RuntimeError(f'{fmt.upper()}encoding is not available on this processing server.') from exc
                finally:
                    image.close()
                generated.append(path)
            if fmt in {'gif', 'tiff'}:
                if not images:
                    raise ValueError('The PDF has no selected pages.')
                first, rest = (images[0], images[1:])
                if fmt == 'gif':
                    first.save(output, 'GIF', save_all=True, append_images=rest, duration=max(100, int(options.get('frame_duration_ms', 700))), loop=0)
                else:
                    first.save(output, 'TIFF', save_all=True, append_images=rest, compression='tiff_deflate')
                return
            with zipfile.ZipFile(output, 'w', zipfile.ZIP_DEFLATED) as archive:
                for path in generated:
                    archive.write(path, arcname=path.name)
        finally:
            for image in images:
                image.close()

def _searchable_from_images(files: list[Path], output: Path, options: dict[str, Any], workdir: Path) -> None:
    language = _recognition_languages(options)
    executable = legacy.command_path('')
    if not executable:
        raise RuntimeError('is not installed on the processing server.')
    parts: list[Path] = []
    page_index = 0
    try:
        for file in files:
            for frame in _iter_input_frames(file):
                page_index += 1
                prepared = _prepare_scan_image(frame, options, preserve_color=True)
                source_path = workdir / f'searchable-source-{page_index:03d}.png'
                part = workdir / f'searchable-page-{page_index:03d}.pdf'
                try:
                    prepared.save(source_path, 'PNG', optimize=True)
                    legacy._processing_engine_searchable_pdf(executable, source_path, part, language, workdir, page_index)
                    parts.append(part)
                finally:
                    frame.close()
                    prepared.close()
                    source_path.unlink(missing_ok=True)
        if not parts:
            raise ValueError('No readable pages were found.')
        merged = fitz.open()
        try:
            for part in parts:
                with fitz.open(part) as page_pdf:
                    merged.insert_pdf(page_pdf)
            merged.save(output, garbage=4, deflate=True, clean=True)
        finally:
            merged.close()
    finally:
        for part in parts:
            part.unlink(missing_ok=True)

def _searchable_from_pdf(source: Path, output: Path, options: dict[str, Any], workdir: Path) -> None:
    dpi = max(150, min(400, int(options.get('dpi', DEFAULT_SCAN_DPI))))
    rendered: list[Path] = []
    try:
        with fitz.open(source) as document:
            for index, page in enumerate(document, start=1):
                image = _render_page(page, dpi)
                path = workdir / f'searchable-input-{index:03d}.png'
                try:
                    image.save(path, 'PNG')
                finally:
                    image.close()
                rendered.append(path)
        _searchable_from_images(rendered, output, options, workdir)
    finally:
        for path in rendered:
            path.unlink(missing_ok=True)

def run_conversion(spec: legacy.ConversionSpec, files: list[Path], output: Path, options: dict[str, Any], workdir: Path, source_url: str | None=None) -> None:
    """Quality-first website conversion entrypoint. The public FastAPI service executes conversions through ``job_worker``. This wrapper upgrades the user-facing /PDF/image paths while retaining the proven legacy processors for formats where the existing implementation is already the safer option. MCP code is intentionally not changed by this layer."""
    legacy.validate_input_files(spec, files)
    available, reason = legacy.tool_available(spec)
    if not available:
        raise RuntimeError(reason or 'This conversion is unavailable on the processing server.')
    processor = spec.processor
    source = files[0] if files else None
    if run_r19_conversion(spec, files, output, options, workdir, source_url):
        legacy.validate_output_file(output, spec.output_extension)
        return
    if processor == 'pdf_text':
        _write_text(_extract_pdf_models(source, options), output)
    elif processor == 'pdf_docx':
        models = _extract_pdf_models(source, options)
        _write_layout_docx(source, models, output, options)
    elif processor == 'pdf_xlsx':
        _write_xlsx_tables(source, output, options)
    elif processor == 'pdf_csv':
        _write_csv_tables(source, output, options)
    elif processor == 'images_to_pdf':
        _images_to_pdf_quality(files, output, options, scan=False)
    elif processor == 'scan_images_pdf':
        _images_to_pdf_quality(files, output, options, scan=True)
    elif processor.startswith('pdf_to_image:'):
        _pdf_to_images_quality(source, output, processor.split(':', 1)[1], options, workdir)
    else:
        legacy.convert(spec, files, output, options, workdir, source_url)
        legacy.validate_output_file(output, spec.output_extension)
        return
    legacy.validate_output_file(output, spec.output_extension)
